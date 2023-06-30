# -*- coding: utf-8 -*-
""" AcademyTestsTest

This module contains the academy.tests an unique Odoo model
which contains all academy tests attributes and behavior.

This model is the representation of the real life academy tests
"""

from logging import getLogger
from operator import itemgetter
from os import linesep
from odoo.exceptions import ValidationError, UserError
from odoo.osv.expression import AND

import docx
from io import BytesIO
import os
import base64
from re import split

# pylint: disable=locally-disabled, E0401
from odoo import models, fields, api
from odoo.tools.translate import _

from odoo.osv.expression import FALSE_DOMAIN
from .utils.sql_operations import ACADEMY_TESTS_SHUFFLE
from .utils.sql_operations import ACADEMY_TESTS_ARRANGE_BLOCKS
from .utils.sql_inverse_searches import QUESTION_COUNT_SEARCH
from .utils.sql_inverse_searches import SEARCH_TEST_ATTEMPT_COUNT
from .utils.libuseful import prepare_text, fix_established, is_numeric, \
    eval_domain

_logger = getLogger(__name__)


# pylint: disable=locally-disabled, R0903, W0212
class AcademyTestsTest(models.Model):
    """ Stored tests which can be reused in future
    """

    _name = 'academy.tests.test'
    _description = u'Academy tests, test'

    _rec_name = 'name'
    _order = 'write_date DESC, create_date DESC'

    _inherit = [
        'academy.abstract.spreadable',
        'ownership.mixin',
        'image.mixin',
        'mail.thread'
    ]

    name = fields.Char(
        string='Name',
        required=True,
        readonly=False,
        index=True,
        default=None,
        help="Name for this test",
        size=255,
        translate=True,
        track_visibility='onchange'
    )

    description = fields.Text(
        string='Description',
        required=False,
        readonly=False,
        index=False,
        default=None,
        help='Something about this test',
        translate=True
    )

    active = fields.Boolean(
        string='Active',
        required=False,
        readonly=False,
        index=False,
        default=True,
        help=('If the active field is set to false, it will allow you to '
              'hide record without removing it')
    )

    code = fields.Char(
        string='Code',
        required=False,
        readonly=False,
        index=True,
        default=None,
        help='Internal code',
        size=20,
        translate=False
    )

    preamble = fields.Text(
        string='Preamble',
        required=False,
        readonly=False,
        index=False,
        default=lambda self: self.default_preamble(),
        help='What it is said before beginning to test',
        translate=True
    )

    @api.model
    def default_preamble(self):
        return _('This exercise poses different questions, presenting a set '
                 'of alternative answers for each of them, among which you '
                 'must select the only correct one.')

    question_ids = fields.One2many(
        string='Questions',
        required=False,
        readonly=False,
        index=True,
        default=None,
        help=False,
        comodel_name='academy.tests.test.question.rel',
        inverse_name='test_id',
        domain=[],
        context={},
        auto_join=False,
        limit=None,
    )

    last_edition = fields.Datetime(
        string='Last edition',
        required=False,
        readonly=True,
        index=False,
        default=fields.datetime.now(),
        help='Last editing date/time',
        compute='_compute_last_edition'
    )

    @api.depends('write_date', 'question_ids')
    def _compute_last_edition(self):
        for record in self:
            path = 'question_ids.write_date'
            links = record.mapped(path)
            path = 'question_ids.question_id.write_date'
            questions = record.mapped(path)
            path = 'question_ids.question_id.answer_ids.write_date'
            answers = record.mapped(path)
            test = [record.write_date]

            dates = test + links + questions + answers

            record.last_edition = max(dates)

    answers_table_ids = fields.One2many(
        string='Answers table',
        required=False,
        readonly=True,
        index=False,
        default=None,
        help='Summary with answers table',
        comodel_name='academy.tests.answers.table',
        inverse_name='test_id',
        domain=[],
        context={},
        auto_join=False,
        limit=None,
    )

    assignment_ids = fields.One2many(
        string='Training assignments',
        required=False,
        readonly=False,
        index=True,
        default=None,
        help='False',
        comodel_name='academy.tests.test.training.assignment',
        inverse_name='test_id',
        domain=[],
        context={},
        auto_join=False,
        limit=None
    )

    assignment_count = fields.Integer(
        string='Nº assignments',
        required=False,
        readonly=True,
        index=False,
        default=0,
        help='Show the number or training assignments for this test',
        store=False,
        compute='_compute_assignment_count'
    )

    @api.depends('assignment_ids')
    def _compute_assignment_count(self):
        for record in self:
            record.assignment_count = \
                len(record.assignment_ids)

    random_template_id = fields.Many2one(
        string='Template',
        required=False,
        readonly=True,
        index=False,
        default=None,
        help='Template has been used to Populate this tests',
        comodel_name='academy.tests.random.template',
        domain=[],
        context={},
        ondelete='cascade',
        auto_join=False
    )

    test_kind_id = fields.Many2one(
        string='Kind of test',
        required=True,
        readonly=False,
        index=False,
        default=lambda self: self.default_test_kind_id(),
        help='Choose the kind for this test',
        comodel_name='academy.tests.test.kind',
        domain=[],
        context={},
        ondelete='cascade',
        auto_join=False
    )

    def default_test_kind_id(self):
        return self.env.ref('academy_tests.academy_tests_test_kind_common')

    first_use_id = fields.Many2one(
        string='First use',
        required=False,
        readonly=False,
        index=False,
        default=None,
        help=False,
        comodel_name='res.partner',
        domain=[],
        context={},
        ondelete='cascade',
        auto_join=False
    )

    authorship = fields.Boolean(
        string='Authorship',
        required=False,
        readonly=False,
        index=False,
        default=True,
        help='Check it to indicate that it is your own authorship'
    )

    repeat_images = fields.Boolean(
        string='Repeat images',
        required=False,
        readonly=False,
        index=False,
        default=False,
        help='Repeat the image every time it is referred to in a question'
    )

    tag_ids = fields.Many2many(
        string='Tags',
        required=False,
        readonly=False,
        index=True,
        default=None,
        help='Tag can be used to better describe this question',
        comodel_name='academy.tests.tag',
        relation='academy_tests_test_tag_rel',
        column1='test_id',
        column2='tag_id',
        domain=[],
        context={},
        limit=None,
        track_visibility='onchange',
    )

    test_block_ids = fields.Many2manyView(
        string='Test blocks',
        required=False,
        readonly=True,
        index=False,
        default=None,
        help='List all blocks have been used in this tests',
        comodel_name='academy.tests.test.block',
        relation='academy_tests_test_test_block_rel',
        column1='test_id',
        column2='test_block_id',
        domain=[],
        context={},
        limit=None,
        copy=False
    )

    auto_arrange_blocks = fields.Boolean(
        string='Auto arrange blocks',
        required=False,
        readonly=False,
        index=False,
        default=True,
        help='Check it to auto arrange questions in blocks'
    )

    restart_numbering = fields.Boolean(
        string='Restart numbering',
        required=False,
        readonly=False,
        index=False,
        default=False,
        help='Check it to restart numbering in each block'
    )

    block_starts_page = fields.Boolean(
        string='Block starts a page',
        required=False,
        readonly=False,
        index=False,
        default=True,
        help='Check it to do each block starts a new page'
    )

    # -------------------------- MANAGEMENT FIELDS ----------------------------

    attempt_count = fields.Integer(
        string='Attempt count',
        required=False,
        readonly=True,
        index=False,
        default=0,
        store=False,
        help='Show number of test attempts',
        compute='_compute_attempt_count',
        search='_search_attempt_count'
    )

    @api.depends('attempt_ids')
    def _compute_attempt_count(self):
        for record in self:
            record.attempt_count = len(record.attempt_ids)

    def _search_attempt_count(self, operator, value):
        domain = FALSE_DOMAIN
        operator, value = self._ensure_search_attempt_count_(operator, value)
        query = SEARCH_TEST_ATTEMPT_COUNT.format(operator, value)

        self.env.cr.execute(query)
        rows = self.env.cr.dictfetchall()

        if rows:
            test_ids = [row['test_id'] for row in rows]
            domain = [('id', 'in', test_ids)]

        return domain

    @staticmethod
    def _ensure_search_attempt_count_(operator, value):
        if isinstance(value, bool):
            if not operator == '=':
                value = not value

            if value:
                operator = '>'

            value = 0

        return operator, value

    question_count = fields.Integer(
        string='Number of questions',
        required=False,
        readonly=False,
        index=False,
        default=0,
        store=False,
        help='Show the number of questions in test',
        compute='_compute_question_count',
        search='_search_question_count'
    )

    @api.depends('question_ids')
    def _compute_question_count(self):
        for record in self:
            record.question_count = len(record.question_ids)

    def _search_question_count(self, operator, operand):
        supported = ['=', '!=', '<=', '<', '>', '>=']

        assert operator in supported, \
            UserError(_('Search operator not supported'))

        assert is_numeric(operand) or operand in [True, False], \
            UserError(_('Search value not supported'))

        operator, operand = fix_established(operator, operand)

        sql = QUESTION_COUNT_SEARCH.format(operator, operand)

        self.env.cr.execute(sql)
        ids = self.env.cr.fetchall()

        return [('id', 'in', ids)]

    topic_ids = fields.Many2manyView(
        string='Topics',
        required=False,
        readonly=True,
        index=False,
        default=None,
        help='Topics from all the questions in the test',
        comodel_name='academy.tests.topic',
        relation='academy_tests_test_topic_rel',
        column1='test_id',
        column2='topic_id',
        domain=[],
        context={},
        limit=None,
        copy=False
    )

    topic_count = fields.Integer(
        string='Number of topics',
        required=False,
        readonly=True,
        index=False,
        default=0,
        store=False,
        help='Display the number of topics related with test',
        compute=lambda self: self._compute_topic_count()
    )

    @api.depends('question_ids')
    def _compute_topic_count(self):
        for record in self:
            question_set = record.question_ids.mapped('question_id')
            topic_set = question_set.mapped('topic_id')
            ids = topic_set.mapped('id')

            record.topic_count = len(ids)

    topic_id = fields.Many2one(
        string='Topic',
        required=False,
        readonly=True,
        index=False,
        default=None,
        help=False,
        comodel_name='academy.tests.topic',
        domain=[],
        context={},
        ondelete='cascade',
        auto_join=False,
        compute=lambda self: self._compute_topic_id()
    )

    @api.depends('question_ids')
    def _compute_topic_id(self):
        for record in self:
            rel_ids = record.question_ids.filtered(
                lambda rel: rel.question_id.topic_id)
            question_ids = rel_ids.mapped('question_id')
            topics = {k.id: 0 for k in question_ids.mapped('topic_id')}

            if not topics:
                record.topic_id = None
            else:
                for question_id in question_ids:
                    _id = question_id.topic_id.id
                    topics[_id] = topics[_id] + 1

                topic_id = max(topics.items(), key=itemgetter(1))[0]

                topic_obj = self.env['academy.tests.topic']
                record.topic_id = topic_obj.browse(topic_id)

    lang = fields.Char(
        string='Language',
        required=True,
        readonly=True,
        index=False,
        help=False,
        size=255,
        translate=False,
        compute='_compute_lang',
        store=False
    )

    attempt_ids = fields.One2many(
        string='Attempts',
        required=False,
        readonly=False,
        index=True,
        default=None,
        help='Related test attempts',
        comodel_name='academy.tests.attempt',
        inverse_name='test_id',
        domain=[],
        context={},
        auto_join=False,
        limit=None
    )

    correction_scale_id = fields.Many2one(
        string='Correction scale',
        required=True,
        readonly=False,
        index=False,
        default=lambda self: self.default_correction_scale_id(),
        help='Choose the scale of correction',
        comodel_name='academy.tests.correction.scale',
        domain=[],
        context={},
        ondelete='cascade',
        auto_join=False
    )

    def default_correction_scale_id(self):
        xid = 'academy_tests.academy_tests_correction_scale_default'
        return self.env.ref(xid)

    # -------------------------- PYTHON CONSTRAINS ----------------------------

    @api.constrains('question_ids')
    def _check_question_availability(self):
        """ Check if questions are ready and they have not dependencies
        """
        dep_msg = _('Some of the questions have unmet dependencies.')
        ready_msg = _('Some of the questions have not been marked as ready.')

        for record in self:
            for link_id in record.question_ids.sorted('sequence'):
                previous_ids = record.question_ids.filtered(
                    lambda x: x.sequence < link_id.sequence) \
                    .mapped('question_id')

                dep_id = link_id.question_id.depends_on_id

                if dep_id and dep_id not in previous_ids:
                    raise ValidationError(dep_msg)

                if link_id.question_id.status != 'ready':
                    raise ValidationError(ready_msg)

    # ----------------------- AUXILIARY FIELD METHODS -------------------------

    @api.depends('name')
    def _compute_lang(self):
        """ Gets the language used by the current user and sets it as `lang`
            field value
        """

        user_id = self.env['res.users'].browse(self.env.uid)

        for record in self:
            record.lang = user_id.lang

    def import_questions(self):
        """ Runs a wizard to import questions from plain text
        @note: actually this method is not used
        """
        return {
            'type': 'ir.actions.act_window',
            'res_model': 'academy.tests.question.import.wizard',
            'view_mode': 'form',
            'views': [(False, 'form')],
            'target': 'new',
            'context': {'default_test_id': self.id}
        }

    def random_questions(self):
        """ Runs wizard to append random questions. This allows uses to set
        filter criteria, maximum number of questions, etc.
        """
        return {
            'type': 'ir.actions.act_window',
            'res_model': 'academy.tests.random.wizard',
            'view_mode': 'form',
            'views': [(False, 'form')],
            'target': 'new',
            'context': {'default_test_id': self.id}
        }

    def show_questions(self):
        """ Runs default view for academy.tests.question with a filter to
        show only current test questions
        """

        self.ensure_one()

        xid = 'academy_tests.action_test_question_links_act_window'
        action = self.env.ref(xid)
        domain = eval_domain(action.domain)
        link_ids = self.mapped('question_ids.id')

        return {
            'name': _('Questions links'),
            'view_mode': 'kanban,tree,form,pivot',
            'res_model': 'academy.tests.test.question.rel',
            'type': 'ir.actions.act_window',
            'target': 'current',
            'domain': AND([domain, [('id', 'in', link_ids)]]),
            'context': {'default_test_id': self.id},
        }

    @api.model
    def create(self, values):
        """ Create a new record for a model AcademyTestsTest
            @param values: provides a data for new record

            @return: returns a id of new record
        """

        result = super(AcademyTestsTest, self).create(values)
        result.resequence()

        return result

    def write(self, values):
        """ Update all record(s) in recordset, with new value comes as {values}
            @param values: dict of new values to be set

            @return: True on success, False otherwise
        """

        result = super(AcademyTestsTest, self).write(values)
        self.resequence()

        return result

    @api.returns('self', lambda value: value.id)
    def copy(self, default=None):
        self.ensure_one()

        # STEP 1: Ensure default is a dictionary
        if default is None:
            default = {}

        # STEP 2: Make new name adding ``(copy)``
        if 'name' not in default:
            default['name'] = _("%s (copy)") % self.name

        # STEP 3: Create new links for all questions in the original test
        create_empty = self.env.context.get('create_empty_test', False)
        if self.question_ids and not create_empty:
            leafs = self.question_ids.mapped(self._mapped_question_ids)
            if(leafs):
                default['question_ids'] = leafs

        # STEP 4: Call parent method
        result = super(AcademyTestsTest, self).copy(default=default)

        return result

    @staticmethod
    def _mapped_question_ids(item):
        return (0, 0, {
            'test_id': item.test_id.id,
            'question_id': item.question_id.id,
            'sequence': item.sequence,
            'active': item.active
        })

    def resequence(self):
        """ This updates the sequence of the questions into the test
        """

        if self:

            if self.auto_arrange_blocks:  # Keep test blocks
                test_ids = self.mapped('id')
                test_ids_str = [str(tid) for tid in test_ids]
                joined = ', '.join(test_ids_str)

                query = ACADEMY_TESTS_ARRANGE_BLOCKS.format(joined)

                self.env.cr.execute(query)

            else:
                for record in self:
                    rel_set = record.question_ids.sorted()

                    index = 1
                    for rel_item in rel_set:
                        rel_item.write({'sequence': index})
                        index = index + 1

    def shuffle(self):
        dep_msg = _('This test has dependent questions, '
                    'it must be sorted manually')

        for record in self:

            if not record.question_ids:
                continue

            link_ids = record.question_ids
            dep_ids = link_ids.mapped('question_id.depends_on_id')

            if dep_ids:
                raise UserError(dep_msg)

            else:
                ids = record.mapped('id')

                # Use complex search to keep images consecutive
                target_ids = ', '.join([str(i) for i in ids])
                sql = ACADEMY_TESTS_SHUFFLE.format(target_ids)

                record.env.cr.execute(sql)
                record.env.cr.commit()
                # record.env.invalidate_all()

    def _creation_subtype(self):
        xid = 'academy_tests.academy_tests_test_created'
        return self.env.ref(xid)

    def _track_subtype(self, init_values):
        self.ensure_one()

        if('active' not in init_values):

            if 'owner_id' in init_values:
                xid = 'academy_tests.academy_tests_test_owned'
            else:
                xid = 'academy_tests.academy_tests_test_written'

            return self.env.ref(xid)

        else:
            _super = super(AcademyTestsTest, self)
            return _super._track_subtype(init_values)

    # ---------------------------- PUBLIC METHODS -----------------------------

    def to_string(self, editable=False):
        """
        @param editable (bool): if it set to true IDs will be preserved,
        otherwise index or URLs will be used instead
        """

        text = ''

        for record in self:
            parts = []

            parts.append('# {}'.format(record.name.strip()))

            desc = prepare_text(record.description or '', '>')
            if(desc):
                parts.append(desc)

            pre = prepare_text(record.preamble or '')
            if(pre):
                parts.append(pre)

            parts.append(linesep)

            for link in self.question_ids:
                question = link.question_id

                parts.append(question.to_string(editable))

            text += linesep.join(parts)

        return text

    def choose_report_dialog(self):
        return {
            'name': _('Choose report'),
            'type': 'ir.actions.act_window',
            'res_model': 'academy.test.choose.report.wizard',
            'view_mode': 'form',
            'views': [(False, 'form')],
            'target': 'new',
            'domain': [],
            'context': {'default_test_id': self.id}
        }

    def update_questions_dialog(self):
        wizard_model = 'academy.tests.update.questions.wizard'
        question_set = self.mapped('question_ids.question_id')

        wizard_set = self.env[wizard_model]
        wizard_set = wizard_set.create({})
        wizard_set.set_questions(question_set)

        return {
            'name': _('Update questions'),
            'type': 'ir.actions.act_window',
            'res_model': wizard_model,
            'view_mode': 'form',
            'views': [(False, 'form')],
            'target': 'new',
            'domain': [],
            'res_id': wizard_set.id
        }

    def save_as_docx(self):

        tpl_path = self._docx_get_template_path()
        wd = self._docx_new_from_template(tpl_path)

        self._docx_update_title(wd)
        self._docx_update_preamble(wd)

        self._docx_update_questions(wd)

        buffer = BytesIO()
        wd.save(buffer)
        datas = base64.b64encode(buffer.getvalue())

        attach_item = self.env['ir.attachment'].create({
            'name': '{}.docx'.format(self.name),
            'datas': datas,
            'type': 'binary',
            'res_model': 'academy.tests.test',
            'res_id': self.id,
            'mimetype': 'application/msword'
        })

        return {
            'name': attach_item.name,
            'res_model': 'ir.actions.act_url',
            'type': 'ir.actions.act_url',
            'target': '_blank',
            'url': '/web/content/{}?download=true'.format(attach_item.id)
        }

    def _docx_get_template_path(self):
        file_path = os.path.realpath(__file__)
        dir_path = os.path.dirname(file_path)
        return os.path.join(dir_path, '..', 'static', 'docx', 'test.docx')

    def _docx_new_from_template(self, template_path):
        wd = None

        with open(template_path, 'rb') as tpl_file:
            buffer = BytesIO(tpl_file.read())
            wd = docx.Document(buffer)

        return wd

    def _docx_update_title(self, wd):
        wd.add_paragraph(self.name, 'Heading 1')

    def _docx_update_preamble(self, wd):
        wd.add_paragraph(self.preamble, 'Preamble')

    @staticmethod
    def _docx_update_block(link, wd):
        if link.test_block_id:
            wd.add_paragraph(link.test_block_id.name, 'Heading 2')
            if link.test_block_id.preamble:
                wd.add_paragraph(link.test_block_id.preamble, 'Preamble')

    def _docx_update_questions(self, wd):
        test_block_id = 0

        for qitem in self.question_ids:
            if test_block_id != qitem.test_block_id.id:
                self._docx_update_block(qitem, wd)

            test_block_id = qitem.test_block_id.id

            for img in qitem.ir_attachment_image_ids:
                content = base64.b64decode(img.datas)
                img_stream = BytesIO(content)
                wd.add_picture(img_stream)
                wd.add_paragraph(img.name, 'Caption')

            if qitem.description:
                for line in self._split_lines(qitem.description):
                    if line:
                        wd.add_paragraph(line, 'About')
            if qitem.preamble:
                for line in self._split_lines(qitem.preamble):
                    if line:
                        wd.add_paragraph(line, 'About')

            wd.add_paragraph(qitem.name, 'Question')

            for aitem in qitem.answer_ids:
                if aitem.is_correct:
                    wd.add_paragraph(style='Answer') \
                        .add_run(aitem.name, style='Right answer')
                else:
                    wd.add_paragraph(aitem.name, 'Answer')

            wd.add_paragraph('', 'Normal')

    @staticmethod
    def _split_lines(content):
        return split(r"[\r\n]+", content)

    def request_for_questions(self):

        self.ensure_one()

        rset_domain = [('test_id', '=', self.id)]
        rset_obj = self.env['academy.tests.question.request.set']
        rset_set = rset_obj.search(rset_domain)

        act_window = {
            'model': 'ir.actions.act_window',
            'type': 'ir.actions.act_window',
            'name': _('Request for questions'),
            'res_model': 'academy.tests.question.request.set',
            'target': 'new',
            'view_mode': 'form',
            'views': [(False, 'form')],
            'domain': [],
            'context': {}
        }

        if rset_set:
            act_window['res_id'] = rset_set.id
        else:
            act_window['context'] = {'default_test_id': self.id}

        return act_window

    def download_as_moodle_xml(self):
        self.ensure_one()

        relative_url = '/academy_tests/moodle/test?test_id={}'
        return {
            'type': 'ir.actions.act_url',
            'url': relative_url.format(self.id),
            'target': 'self',
        }

    def download_as_pdf(self):
        self.ensure_one()

        report_xid = 'academy_tests.action_report_full_printable_test'
        report_act = self.env.ref(report_xid)
        doc_ids = [self.id]

        return report_act.report_action(docids=doc_ids, data={}, config=False)

    def redirect_to_preview(self):
        self.ensure_one()

        relative_url = '/academy_tests/test/preview?test_id={}'
        return {
            'type': 'ir.actions.act_url',
            'url': relative_url.format(self.id),
            'target': 'new',
        }

    def compute_block_classes(self, block):
        self.ensure_one()

        classes = ['academy-post-test-block']

        if self.restart_numbering:
            classes.append('academy-post-test-restart-numbering')

        if self.block_starts_page:
            classes.append('academy-post-test-page-break')

        if block or self.block_starts_page:
            classes.extend(['invisible', 'm-0', 'border-0'])

        return ' '.join(classes)

    def view_training_assignments(self):
        self.ensure_one()

        return {
            'model': 'ir.actions.act_window',
            'type': 'ir.actions.act_window',
            'name': _('Training assignments'),
            'res_model': 'academy.tests.test.training.assignment',
            'target': 'current',
            'view_mode': 'kanban,tree,form',
            'domain': [('test_id', '=', self.id)],
            'context': {
                'name_get': 'training',
                'default_test_id': self.id
            }
        }

    def new_from_template(self, gui=True):
        self.ensure_one()

        if not self.random_template_id:
            msg = _('This test was not created from a template')
            raise UserError(msg)

        result = self.random_template_id.new_test(gui=gui)
        if isinstance(result, dict) and 'target' in result.keys():
            result['target'] = 'main'

        return result

    def new_assignment_to_training(self):
        self.ensure_one()

        return {
            'model': 'ir.actions.act_window',
            'type': 'ir.actions.act_window',
            'name': _('New assignment'),
            'res_model': 'academy.tests.test.training.assignment',
            'target': 'new',
            'view_mode': 'form',
            'context': {
                'default_test_id': self.id
            }
        }
